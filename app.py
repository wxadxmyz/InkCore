# -*- coding: utf-8 -*-
"""
墨核 AI Studio —— 桌面端全能 AI 助手（可打包为 exe）
主界面：对话。全能桌面 AI 助手：聊天 / 问答 / 联网搜索 / 写作 / AI 生图 / 提示词 / 记忆 / Skill / 终端 / 代码 / Agent。

运行（开发）： python3 app.py            → http://127.0.0.1:7860
打包桌面端：   python3 desktop_app.py    → 原生窗口（PyInstaller 打包为 exe，见 build.spec）
"""
import os, re, json, random, uuid, subprocess, sys, io, base64, datetime, urllib.parse, zipfile, shutil
from pathlib import Path
import requests
from flask import Flask, request, jsonify, send_from_directory, send_file, Response, stream_with_context
from docx import Document
import local_llm  # 本地内置模型（离线推理，无需 Ollama）

# ---------- 路径（兼容打包后的 exe） ----------
FROZEN = getattr(sys, "frozen", False)
if FROZEN:
    # PyInstaller 单文件模式：资源被解压到临时目录 _MEIPASS
    APP_DIR = Path(sys._MEIPASS)
else:
    APP_DIR = Path(__file__).resolve().parent
DATA_HOME = Path(os.path.expanduser("~")) / "MoHeAI"          # 用户可写数据目录
DATA_HOME.mkdir(exist_ok=True)
OUT = DATA_HOME / "output"; OUT.mkdir(exist_ok=True)
STATIC = APP_DIR / "static"

CONFIG_PATH = DATA_HOME / "config.json"
# 内置供应商预设（一键切换，免填 base_url）
PROVIDERS = {
    "openai":  {"label":"OpenAI","base_url":"https://api.openai.com/v1","model":"gpt-4o-mini",
                "models":["gpt-4o-mini","gpt-4o","gpt-4o-latest","o1","o1-mini","o3-mini"]},
    "deepseek":{"label":"DeepSeek","base_url":"https://api.deepseek.com/v1","model":"deepseek-chat",
                "models":["deepseek-chat","deepseek-reasoner","deepseek-coder"]},
    "qwen":    {"label":"通义千问","base_url":"https://dashscope.aliyuncs.com/compatible-mode/v1","model":"qwen-plus",
                "models":["qwen-plus","qwen-turbo","qwen-max","qwen-coder-plus"]},
    "moonshot":{"label":"Moonshot","base_url":"https://api.moonshot.cn/v1","model":"moonshot-v1-8k",
                "models":["moonshot-v1-8k","moonshot-v1-32k","moonshot-v1-128k"]},
    "glm":     {"label":"智谱 GLM","base_url":"https://open.bigmodel.cn/api/paas/v4","model":"glm-4-flash",
                "models":["glm-4-flash","glm-4","glm-4-plus","glm-4-air"]},
    "ollama":  {"label":"Ollama(本地)","base_url":"http://localhost:11434/v1","model":"llama3",
                "models":["llama3","qwen2.5","deepseek-r1"]},
    "custom":  {"label":"自定义","base_url":"https://api.openai.com/v1","model":"gpt-4o-mini",
                "models":["gpt-4o-mini"]},
    "siliconflow":{"label":"SiliconFlow(免费聚合)","base_url":"https://api.siliconflow.cn/v1","model":"Qwen/Qwen2.5-7B-Instruct",
                "models":["Qwen/Qwen2.5-7B-Instruct","deepseek-ai/DeepSeek-V3","deepseek-ai/DeepSeek-R1",
                          "Pro/deepseek-ai/DeepSeek-R1","Qwen/Qwen2.5-Coder-7B-Instruct","Qwen/Qwen2.5-VL-7B-Instruct"]},
    "openrouter":{"label":"OpenRouter(免费聚合)","base_url":"https://openrouter.ai/api/v1","model":"openai/gpt-4o-mini",
                "models":["openai/gpt-4o-mini","anthropic/claude-3.5-sonnet","google/gemini-2.0-flash-exp",
                          "deepseek/deepseek-r1","meta-llama/llama-3.3-70b-instruct","mistralai/mixtral-8x7b-instruct"]},
    "embedded":{"label":"本地内置模型(离线·免Ollama)","base_url":"embedded://local","model":"qwen2.5-1.5b-instruct-q4_k_m",
                "models":["qwen2.5-1.5b-instruct-q4_k_m"]},
}
DEFAULT_CONFIG = {
    "active_profile": "default",
    "provider": "custom",
    "llm": {"enabled": False, "base_url": "https://api.openai.com/v1",
            "api_key": "", "model": "gpt-4o-mini", "temperature": 0.9},
    "image": {"enabled": False, "base_url": "https://api.openai.com/v1",
              "api_key": "", "model": "gpt-image-1"},
    "vision": {"enabled": False, "base_url": "https://api.openai.com/v1",
               "api_key": "", "model": "gpt-4o-mini"},
    "sandbox": {"readonly": False},
    "provider_keys": {},   # 各供应商各自记住的 key/base_url/model，切换不丢失
    "hotkey": "ctrl+alt+m",   # 全局热键（桌面端唤起/隐藏窗口），兼容 keyboard 库格式
}
APP_VERSION = "1.0.0"   # 自动更新比对用
UPDATE_URL_DEFAULT = "https://raw.githubusercontent.com/wxadxmyz/InkCore/main/version.json"  # 自动更新清单（替换为你的 GitHub 用户名；也可托管到自己的域名）
def _load(path, default):
    try:
        if path.exists(): return json.loads(path.read_text(encoding="utf-8"))
    except Exception: pass
    return default
def _save(path, obj):
    path.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")

config = _load(CONFIG_PATH, DEFAULT_CONFIG)
def save_config(): _save(CONFIG_PATH, config)

def profile_dir(name=None):
    name = name or config.get("active_profile", "default")
    d = DATA_HOME / "profiles" / name; d.mkdir(parents=True, exist_ok=True)
    return d

# ---------- 默认 Skill ----------
DEFAULT_SKILLS = [
    {"id":"skill_naming","name":"起名器","icon":"tag","desc":"为人物、门派、功法、公司等生成有韵味的名字。",
     "trigger":["起名","取名","名字"],"type":"prompt","content":"你是一位精通网文命名美学的起名师。请根据用户需求，给出 8 个风格统一、朗朗上口且各有寓意的名字，并简要说明每个名字的出处与气质。"},
    {"id":"skill_prompt_img","name":"生图提示词","icon":"image","desc":"把模糊想法转成高质量 AI 生图提示词（含负向提示）。",
     "trigger":["提示词","生图词","prompt"],"type":"prompt","content":"你是一位 Midjourney/SD 提示词工程师。请把用户的想法扩写成：主体 + 风格 + 光线 + 构图 + 细节 + 参数，并附中文负向提示词。"},
]

# ---------- 技能市场（可一键添加的精选技能，区别于默认内置） ----------
SKILL_GALLERY = [
    {"id":"mkt_abstract","name":"文章摘要","icon":"list","desc":"把长文/长对话浓缩成 3 条要点 + 一句话结论。",
     "trigger":["摘要","总结","归纳","提炼"],"type":"prompt","content":"你是一位信息压缩专家。请把用户给出的内容提炼为：① 3 条核心要点（每条≤25字）② 一句话结论。保持原意，不编造。"},
    {"id":"mkt_title","name":"爆款标题","icon":"prompt","desc":"给文章/视频生成 10 个高点击率标题（含数字/悬念/痛点）。",
     "trigger":["标题","爆款","取名","题目"],"type":"prompt","content":"你是一位资深新媒体编辑。请基于用户主题，给出 10 个高点击率标题，覆盖：数字盘点型、悬念型、痛点型、身份认同型。只列标题。"},
    {"id":"mkt_outline","name":"大纲生成","icon":"doc","desc":"把一句话想法扩成文章/小说大纲（起承转合 + 章节）。",
     "trigger":["大纲","框架","结构","提纲"],"type":"prompt","content":"你是一位写作教练。请基于用户的一句话主题，生成完整大纲：背景设定、核心冲突、起承转合四幕、3-5 个章节标题与每章要点。"},
    {"id":"mkt_rewrite","name":"润色改写","icon":"edit","desc":"把生硬文字改写成更有节奏感、更有网感的版本。",
     "trigger":["润色","改写","优化","通顺"],"type":"prompt","content":"你是一位文字润色师。请在不改变原意的前提下，把用户的文字改写得更通顺、更有节奏感和网感，保留关键信息。"},
    {"id":"mkt_translator","name":"本地化翻译","icon":"translate","desc":"中英互译并给出地道表达与术语说明。",
     "trigger":["本地化","地道","术语","翻译"],"type":"prompt","content":"你是一位本地化翻译专家。请先给出准确译文，再列出 2-3 个地道替代表达，并解释关键术语的处理。"},
    {"id":"mkt_brainstorm","name":"头脑风暴","icon":"bulb","desc":"围绕一个话题发散 20 个创意点子。",
     "trigger":["脑暴","创意","点子","灵感"],"type":"prompt","content":"你是一位创意总监。请围绕用户的话题，发散出 20 个不重样、可落地的创意点子，按『低成本/中等/大胆』三档分组。"},
]

# ---------- 离谱玩法（用户可自建 / 导入导出的趣味玩法） ----------
DEFAULT_FUNS = [
    {"id":"fun_clean","name":"清理电脑","emoji":"🧹","desc":"只读扫描磁盘占用","prompt":"清理电脑"},
    {"id":"fun_game","name":"搓游戏","emoji":"🎮","desc":"现场生成可玩 HTML","prompt":"搓一个贪吃蛇"},
    {"id":"fun_bili","name":"连外部站","emoji":"🌐","desc":"终端 curl 探活","prompt":"连 b站"},
    {"id":"fun_dl","name":"下载软件","emoji":"⬇️","desc":"调用终端拉取文件","prompt":"下载 https://example.com/app.zip"},
    {"id":"fun_joke","name":"讲笑话","emoji":"😄","desc":"随机轻松一下","prompt":"讲一个程序员笑话"},
    {"id":"fun_emoji","name":"emoji 剧情","emoji":"🎭","desc":"创意小玩法","prompt":"用emoji写一段武侠剧情"},
]

# ---------- 当前 profile 的内存状态 ----------
memory = {}; conversations = {}; personas = {}; skills = []; funs = []
def load_profile():
    global memory, conversations, personas, skills, funs
    d = profile_dir()
    memory = _load(d/"memory.json", {})
    conversations = _load(d/"conversations.json", {})
    personas = _load(d/"personas.json", {})
    if (d/"skills.json").exists():
        skills = _load(d/"skills.json", [])
    else:
        skills = [dict(s) for s in DEFAULT_SKILLS]; _save(d/"skills.json", skills)
    if (d/"funs.json").exists():
        funs = _load(d/"funs.json", [])
    else:
        funs = [dict(f) for f in DEFAULT_FUNS]; _save(d/"funs.json", funs)
load_profile()

# ---------- 会话级产物清单（改动12：按会话记录 AI 生成的图片/游戏/HTML 等） ----------
ARTIFACTS_PATH = DATA_HOME / "artifacts.json"
def load_artifacts():
    return _load(ARTIFACTS_PATH, {})
def record_artifact(session_id, art):
    session_id = session_id or "default"
    a = load_artifacts()
    a.setdefault(session_id, []).append(art)
    _save(ARTIFACTS_PATH, a)

# ---------- 记忆系统 ----------
def get_mem(session_id):
    return memory.setdefault(session_id, {"facts":[],"characters":[],"world":[],"style":"","projects":[]})
def update_memory(session_id, text):
    m = get_mem(session_id)
    for pat in [r"主角[叫是](\S{1,6})", r"(\S{1,6})(?:是|饰演)(男主|女主|反派|主角)", r"人物[：:]\s*(\S{1,8})", r"品牌[：:]\s*(\S{1,8})"]:
        for mm in re.findall(pat, text):
            name = mm if isinstance(mm, str) else mm[0]
            if name and name not in m["characters"]:
                m["characters"].append(name)
    for kw in ["仙侠","都市","悬疑","科幻","言情","玄幻","历史","末日","西幻","无限流"]:
        if kw in text and f"偏好题材：{kw}" not in m["facts"]:
            m["facts"].append(f"偏好题材：{kw}")
    sm = re.search(r"风格[：: ]*(?:是|要|喜欢|偏向|为)?\s*([^\s，。,；;：:]{1,10})", text)
    if sm:
        st = sm.group(1).lstrip("是").strip()
        if st: m["style"] = st
    _save(profile_dir()/"memory.json", memory)
def memory_context(session_id):
    m = get_mem(session_id); bits=[]
    if m["characters"]: bits.append("人物："+"、".join(m["characters"]))
    if m["style"]: bits.append("风格："+m["style"])
    if m["facts"]: bits.append("设定："+"。".join(m["facts"]))
    if m["world"]: bits.append("世界观："+"、".join(m["world"]))
    return ("（已知创作设定："+"。".join(bits)+"）\n") if bits else ""

# ---------- 真实大模型 ----------
SYSTEM_BASE = "你是「墨核 AI」桌面智能助手，能自由聊天、回答各类问题、联网搜资料、写作（小说/文章/文案）、生成提示词、写代码、给建议。风格自然、准确、有帮助、信息密度高。"
SYSTEM_WRITING = SYSTEM_BASE + "根据用户需求创作：短篇/小说/漫剧/短剧/剧本/公众号文章等。输出结构完整、可直接使用的中文内容；若用户给出主角名/题材/风格请严格遵循。"
SYSTEM_PROMPT = SYSTEM_BASE + "你是 Midjourney/SD 提示词工程师，把用户想法扩写成：主体+风格+光线+构图+细节+参数，并附中文负向提示词。给出英文正向 + 中文负向。"
SYSTEM_SEARCH = SYSTEM_BASE + "你具备联网搜索能力。请结合用户提供的搜索结果（标注 [搜索]）作答，并注明信息要点；若结果不足，明确说明。"
SYSTEM_TRANSLATE = "你是专业中英互译引擎。仅输出译文，保留原文格式与专有名词；若原文为中文则译为英文，若为英文则译为中文。"
SYSTEM_CLASSIFY = ('你是意图分类器。仅回复一个 JSON：{"intent":"writing|image|prompt|terminal|code|memory|chat|search|agent|vision|translate|fun|skill","need_clarify":true/false}。'
                  'intent 取最匹配的一项；若用户需求模糊、无法判断要做什么，need_clarify 为 true。不要输出其他内容。')

def get_llm_cfg(session_id=None):
    """返回某会话生效的 LLM 配置：若该会话单独设置了 model，则在全局配置基础上覆盖模型名。"""
    cfg = dict(config["llm"])
    if session_id:
        conv = conversations.get(session_id)
        m = (conv or {}).get("model")
        if m:
            cfg["model"] = m
    return cfg

def _auth_required(c):
    """本地模型（Ollama 等）无需 API Key；远程厂商需要。用于放开「零 Key 本地推理」路径。"""
    bu = (c.get("base_url") or "").lower()
    if "ollama" in bu or bu.startswith("http://localhost") or bu.startswith("http://127.") or bu.startswith("http://0.0.0.0") or bu.startswith("embedded://"):
        return False
    return True

def _llm_usable(cfg):
    """判断某会话是否真的能用 LLM：需启用，且有 Key 或指向本地（无需 Key）模型。"""
    if not cfg.get("enabled"): return False
    return bool(cfg.get("api_key")) or not _auth_required(cfg)

def _stream_completion(url, api_key, msgs, model, temperature):
    """返回一个生成器，逐块产出 LLM 的回答文本（SSE / stream=True）。"""
    def gen():
        headers = {"Authorization": f"Bearer {api_key}"} if api_key else {}
        try:
            r = requests.post(url, headers=headers,
                              json={"model": model, "messages": msgs,
                                    "temperature": temperature, "stream": True},
                              timeout=120, stream=True)
            r.raise_for_status()
            for line in r.iter_lines():
                if not line: continue
                s = line.decode("utf-8", "ignore")
                if not s.startswith("data:"): continue
                payload = s[5:].strip()
                if payload == "[DONE]": break
                try:
                    j = json.loads(payload)
                except Exception:
                    continue
                ch = j.get("choices", [{}])[0].get("delta", {}).get("content") or ""
                if ch:
                    yield ch
        except Exception as e:
            yield f"\n\n[LLM 调用失败] {e} —— 已回退到本地模板。"
    return gen()

def _llm_stream(messages, c):
    """统一流式入口：内置本地模型走 llama-cpp-python（进程内），其余走 OpenAI 兼容 HTTP。"""
    bu = (c.get("base_url") or "")
    if bu.startswith("embedded://"):
        return local_llm.stream(messages, c.get("model"), c.get("temperature", 0.9))
    url = bu.rstrip("/") + "/chat/completions"
    return _stream_completion(url, c.get("api_key", ""), messages, c.get("model", "gpt-4o-mini"), c.get("temperature", 0.9))

def _llm_complete(messages, c):
    """统一非流式入口（用于本地回退判断等）。"""
    bu = (c.get("base_url") or "")
    if bu.startswith("embedded://"):
        try:
            return "".join(local_llm.stream(messages, c.get("model"), c.get("temperature", 0.9)))
        except Exception as e:
            return f"[LLM 调用失败] {e} —— 已回退到本地模板。"
    payload = {"model": c.get("model", "gpt-4o-mini"), "messages": messages,
               "temperature": c.get("temperature", 0.9)}
    try:
        r = requests.post(bu.rstrip("/") + "/chat/completions",
                          headers={"Authorization": f"Bearer {c.get('api_key','')}"},
                          json=payload, timeout=60)
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"[LLM 调用失败] {e} —— 已回退到本地模板。"

def call_api(section, system, history, user, image_b64=None, cfg=None, stream=False):
    """通用 OpenAI 兼容调用；section 可为 llm / vision。cfg 可传入覆盖（如按会话选模型）。
    stream=True 时返回生成器（逐块文本），否则返回完整字符串。"""
    c = cfg or (config.get(section) or config["llm"])
    if not c.get("enabled") or (not c.get("api_key") and _auth_required(c)):
        return None
    url = c["base_url"].rstrip("/") + "/chat/completions"
    msgs = [{"role":"system","content":system}]
    for h in (history or [])[-12:]:
        msgs.append({"role": h.get("role"), "content": h.get("content","")})
    content = user
    if image_b64:
        content = [{"type":"text","text":user},
                   {"type":"image_url","image_url":{"url":f"data:image/png;base64,{image_b64}"}}]
    msgs.append({"role":"user","content":content})
    if stream:
        return _llm_stream(msgs, c)
    return _llm_complete(msgs, c)


def call_llm(system, history, user, cfg=None, stream=False):
    c = cfg or config["llm"]
    if not c.get("enabled") or (not c.get("api_key") and _auth_required(c)):
        return None
    url = c["base_url"].rstrip("/") + "/chat/completions"
    msgs = [{"role":"system","content":system}]
    for h in (history or [])[-12:]:
        msgs.append({"role": h.get("role"), "content": h.get("content","")})
    msgs.append({"role":"user","content":user})
    if stream:
        return _llm_stream(msgs, c)
    return _llm_complete(msgs, c)

def _wrap_stream(d, mode, **extra):
    """把 llm_or 返回的 {reply} / {reply_stream} 统一成 route() 的返回结构。"""
    if "reply_stream" in d:
        return {"reply_stream": d["reply_stream"], "mode": mode, **extra}
    return {"reply": d["reply"], "mode": mode, **extra}

def llm_or(system, user, template_fn, session_id, history, stream=False):
    cfg = get_llm_cfg(session_id)
    if _llm_usable(cfg):
        if stream:
            return {"reply_stream": call_llm(system, history, user, cfg=cfg, stream=True)}
        r = call_llm(system, history, user, cfg=cfg)
        if r and not r.startswith("[LLM"):
            return {"reply": r}
    return {"reply": template_fn()}

# ---------- 写作引擎（本地模板，LLM 关闭时或回退用） ----------
GENRE_TEMPLATES = {"短篇":"短篇故事","小说":"小说","漫剧":"漫剧分镜","短剧":"短剧脚本",
                   "剧本":"剧本","脚本":"脚本","公众号":"公众号文章","文章":"文章",
                   "随笔":"文章","文案":"文案","干货":"文章","软文":"文案"}
# 题材词：从“用户需求”里剔除，避免被误当成作品标题（如“写一篇仙侠小说”→ 标题不应该是“仙侠”）
GENRE_WORDS = "仙侠|都市|悬疑|科幻|言情|玄幻|历史|末日|西幻|无限流|爽文|古言|现言|甜宠|虐恋|权谋|复仇|末世|穿书|重生|古风|现代|校园|职场|灵异|盗墓|修真|武侠|军事|种田|系统|直播|美食|游戏"
def detect_genre(text):
    for k,v in GENRE_TEMPLATES.items():
        if k in text: return k,v
    return None
def gen_writing(text, session_id):
    m = get_mem(session_id)
    genre_key, genre_val = detect_genre(text) or ("文章","文章")
    char = None
    cm = re.search(r"主角[叫是](\S{1,6})", text) or re.search(r"(\S{1,6})(?:是|饰演)(男主|女主|反派|主角)", text)
    char = cm.group(1) if cm else (m["characters"][0] if m["characters"] else None)
    chars = char or "林默"
    topic = re.sub(r"(写一篇|一篇|写篇|帮我写|帮我|创作|生成|来一篇|写个|写一段|来一段|一段|短篇|小说|漫剧|短剧|剧本|脚本|公众号|文章|随笔|文案|干货|软文|产品|带货|为|关于|的|主角叫\S+|，|,|"+GENRE_WORDS+")", " ", text)
    topic = re.sub(r"\s+"," ",topic).strip() or (char or "未命名作品"); topic = topic[:16]
    style = m["style"] or "快节奏、强钩子"
    parts=[]
    if genre_key in ("漫剧","短剧","剧本","脚本"):
        parts.append(f"# 🎬 《{topic}》{genre_val}（3 分钟样稿）")
        parts.append(f"> 风格：{style}　|　主要人物：{chars}")
        for name,shot,desc in [
            ("场景1·开场钩子","近景·特写",f"{chars}在雨夜的旧公寓接到一通神秘电话，对方只说了一句：『你父亲不是死于意外。』"),
            ("场景2·转折","中景·跟拍",f"{chars}翻出尘封的日记，发现一个被抹去的名字，决定深夜潜入公司档案室。"),
            ("场景3·高潮","全景·手持",f"档案室火警骤响，黑影逼近，{chars}在电梯即将关闭的瞬间扑出，手中紧攥一份证据。"),
            ("场景4·收束钩子","远景·升镜",f"镜头拉远，城市灯火中{chars}拨通报警电话——但电话那头，是熟悉的声音。")]:
            parts.append(f"\n**{name}**　〔{shot}〕\n　台词/旁白：{desc}")
    elif genre_key=="公众号":
        parts.append(f"# {topic}：这 3 件事，越早明白越好")
        parts.append("\n> 引言：在这个信息过载的时代，我们缺的不是知识，而是判断力。今天聊点扎心的。\n")
        for i,h in enumerate(["第一，别把忙碌当成长","第二，关系比资源更保值","第三，长期主义才是护城河"],1):
            parts.append(f"\n## {h}\n很多人以为{['加班到深夜就是努力','微信好友越多越牛','追风口就能翻身'][i-1]}，其实恰恰相反。真正的{['成长','人脉','机会']}，都藏在那些『不被看见』的坚持里。")
        parts.append("\n## 写在最后\n与君共勉。如果这篇对你有用，点个『在看』，我们下篇见。")
    elif genre_key=="文案":
        parts.append(f"# ✍️ 文案：《{topic}》")
        parts.append("\n**标题备选：**")
        for t in [f"震惊！{topic}原来还能这样玩", f"别再瞎忙了，{topic}的正确打开方式", f"关于{topic}，90% 的人都搞错了"]:
            parts.append(f"- {t}")
        parts.append(f"\n**正文：**\n{topic}，其实没那么复杂。抓住用户最关心的痛点，用一句话说清价值，再给一个无法拒绝的行动指令，转化率自然就上来了。")
        parts.append("\n**行动号召（CTA）：** 点击下方，立即体验 ▸")
    elif genre_key in ("小说","短篇"):
        parts.append(f"# 《{topic}》· {genre_val}开篇")
        parts.append(f"> 设定：{style}　|　人物：{chars}")
        parts.append(f"\n**【楔子】**\n{chars}从没想过，平凡的日子会在一封没有署名的信里拐弯——信上只写了关于「{topic}」的四个字，还有一句：『别回头。』")
        parts.append(f"\n**【第一章·误入】**\n{chars}循着那四个字走进巷子深处，却在拐角撞见另一个也在找「{topic}」的人。两人的命运，从这一刻被系在一起。")
        parts.append(f"\n**【钩子】**\n当{chars}终于看清信纸背面的落款，呼吸一滞——那是{chars}自己很多年后的笔迹。")
    else:
        parts.append(f"# {topic}")
        parts.append(f"\n> 风格：{style}")
        parts.append("\n## 一、背景与现状\n围绕「%s」，我们先看清楚它为什么重要，以及大多数人的认知盲区在哪里。"%topic)
        parts.append("\n## 二、核心方法\n1. **找准切入点**：从最具体的场景入手，避免空谈。\n2. **建立标准**：用可衡量的指标替代模糊感觉。\n3. **小步快跑**：先做出最小可用版本，再迭代。")
        parts.append("\n## 三、避坑指南\n- 不要贪多，先在一个点打透。\n- 不要追求完美，完成好过完满。")
        parts.append("\n## 写在最后\n方法千万条，行动第一条。今晚就挑一个小点试起来。")
    parts.append("\n\n_——已记入记忆库，可继续让我扩写某一段或改换风格。_")
    return "\n".join(parts)

def gen_prompt(text):
    subj = re.sub(r"(提示词|生图词|prompt|生成|写|给|一张|的|画)", "", text).strip() or "一位古风少女"
    return f"""# 🎨 AI 生图提示词（可直接用于 Midjourney / Stable Diffusion）

**英文正向：**
> {subj}, cinematic lighting, highly detailed, 8k, intricate costume, volumetric light, soft bokeh, artstation trending, masterpiece, ultra wide shot --ar 3:4 --niji 6

**中文负向（Negative）：**
> 低分辨率, 畸形手指, 多余肢体, 模糊, 水印, 文字错位, 过度曝光, 丑陋

**结构拆解：**
- 主体：{subj}
- 风格：电影感 / 二次元（niji）
- 光线：体积光 + 柔焦
- 构图：超宽幅，留白压抑感

_在对话里输入「/image {subj}」可直接生成封面。_"""

# ---------- AI 生图（真实绘图模型 or SVG 回退） ----------
GENRE_BG = {"仙侠":("#2b5876","#4e4376","⛰️"),"玄幻":("#0f2027","#2c5364","🐉"),
            "都市":("#232526","#414345","🏙️"),"科幻":("#000428","#004e92","🪐"),
            "言情":("#ff9a9e","#fecfef","🌸"),"悬疑":("#1f1c2c","#928dab","🔍"),
            "末日":("#3a1c1c","#603813","☣️"),"历史":("#3e5151","#decba4","📜")}
# ---------- 聊天驱动的自动生成（改动3：说一句话 → 生成并自动填入对应面板） ----------
def _strip_code(raw):
    """从 LLM 输出里抽取纯代码（去掉 ``` 围栏）。"""
    m = re.search(r"```(?:python|bash|sh|shell)?\s*(.*?)```", raw, re.DOTALL)
    if m: return m.group(1).strip()
    return raw.strip()
def gen_code(prompt):
    """根据用户需求生成 Python 代码（供代码面板自动填入）。LLM 不可用时返回 None。"""
    cfg = get_llm_cfg("code_auto")
    if not _llm_usable(cfg): return None
    sys = ("你是一位 Python 编程助手。请只输出代码，用 ```python 代码块包裹，不要任何解释与前后赘述。"
           "代码应可直接运行，优先使用标准库，需 pip 的包要先 import。")
    try:
        raw = call_llm(sys, None, "需求："+prompt, cfg=cfg)
    except Exception:
        return None
    if not raw or raw.startswith("[LLM"): return None
    return _strip_code(raw)
def gen_terminal_cmd(prompt):
    """根据用户需求生成一条安全的只读 shell 命令（供终端面板自动填入）。"""
    cfg = get_llm_cfg("term_auto")
    if not _llm_usable(cfg): return None
    sys = ("你是命令行助手。把用户需求转成一条最安全、只读的 shell 命令（如 ls / df -h / du -sh）。"
           "禁止 rm/dd/mkfs/shutdown 等破坏性或写操作。只输出命令本身，不要解释。")
    try:
        raw = call_llm(sys, None, "需求："+prompt, cfg=cfg)
    except Exception:
        return None
    if not raw or raw.startswith("[LLM"): return None
    cmd = _strip_code(raw).splitlines()[0].strip() if _strip_code(raw) else raw.strip().splitlines()[0].strip()
    return cmd or None
def gen_agent_plan(prompt):
    """把用户目标拆成目标 + 1~4 步可执行步骤（供 Agent 面板自动填入）。"""
    cfg = get_llm_cfg("agent_auto")
    if not _llm_usable(cfg): return None
    sys = ('你是任务拆解助手。把用户目标拆成 1-4 个可执行步骤。只输出 JSON：'
           '{"goal":"一句话目标","steps":["步骤1","步骤2"]}，不要解释。')
    try:
        raw = call_llm(sys, None, "目标："+prompt, cfg=cfg)
    except Exception:
        return None
    if not raw or raw.startswith("[LLM"): return None
    m = re.search(r"\{.*\}", raw, re.DOTALL)
    if not m: return None
    try:
        d = json.loads(m.group(0)); return {"goal":d.get("goal",prompt), "steps":d.get("steps",[])}
    except Exception:
        return None
def _extract_title(prompt):
    t = re.sub(r"(画|生成|一张|封面|的|图片|生图|/image|提示词|给)", " ", prompt)
    t = re.sub(r"\s+"," ",t).strip() or "未命名作品"
    return t[:14]
def _build_svg(prompt, title):
    bg = ("#2b5876","#4e4376","✨")
    for k,v in GENRE_BG.items():
        if k in prompt: bg=v; break
    c1,c2,emoji = bg; sid=uuid.uuid4().hex[:8]
    return f'''<svg xmlns="http://www.w3.org/2000/svg" width="600" height="800" viewBox="0 0 600 800">
  <defs>
    <linearGradient id="g{sid}" x1="0" y1="0" x2="1" y2="1"><stop offset="0%" stop-color="{c1}"/><stop offset="100%" stop-color="{c2}"/></linearGradient>
    <radialGradient id="r{sid}" cx="50%" cy="30%" r="80%"><stop offset="0%" stop-color="#ffffff" stop-opacity="0.30"/><stop offset="100%" stop-color="#ffffff" stop-opacity="0"/></radialGradient>
    <radialGradient id="glow{sid}" cx="50%" cy="50%" r="50%"><stop offset="0%" stop-color="#ffffff" stop-opacity="0.40"/><stop offset="100%" stop-color="#ffffff" stop-opacity="0"/></radialGradient>
  </defs>
  <rect width="600" height="800" rx="28" fill="url(#g{sid})"/>
  <rect width="600" height="800" rx="28" fill="url(#r{sid})"/>
  <circle cx="300" cy="248" r="155" fill="url(#glow{sid})"/>
  <text x="300" y="300" font-size="118" text-anchor="middle" opacity="0.96">{emoji}</text>
  <rect x="250" y="372" width="100" height="4" rx="2" fill="#ffffffcc"/>
  <text x="300" y="540" font-size="42" font-family="'Noto Serif SC','Songti SC','SimSun',serif" fill="#fff" text-anchor="middle" font-weight="bold">{title}</text>
  <text x="300" y="602" font-size="18" fill="#ffffffcc" text-anchor="middle" font-style="italic">by 墨核 AI Studio (InkCore)</text>
  <rect x="170" y="628" width="260" height="2" fill="#ffffff55"/>
  <text x="300" y="676" font-size="14" fill="#ffffffaa" text-anchor="middle" letter-spacing="3">AI GENERATED COVER</text>
  <rect x="16" y="16" width="568" height="768" rx="20" fill="none" stroke="#ffffff" stroke-opacity="0.22" stroke-width="2"/>
</svg>'''
def gen_image(prompt, session_id):
    title = _extract_title(prompt)
    c = config["image"]
    path=None; data=None; fmt=None
    if c.get("enabled") and c.get("api_key"):
        try:
            url = c["base_url"].rstrip("/") + "/images/generations"
            r = requests.post(url, headers={"Authorization": f"Bearer {c['api_key']}"},
                              json={"model": c.get("model","gpt-image-1"), "prompt": prompt[:1000], "n":1, "size":"1024x1024"}, timeout=90)
            r.raise_for_status(); item = r.json()["data"][0]
            if item.get("b64_json"):
                raw = base64.b64decode(item["b64_json"])
                fn = OUT/f"cover_{uuid.uuid4().hex[:8]}.png"; fn.write_bytes(raw)
                path=f"/output/{fn.name}"; fmt="png"; data="data:image/png;base64,"+item["b64_json"]
            elif item.get("url"):
                raw = requests.get(item["url"], timeout=60).content
                fn = OUT/f"cover_{uuid.uuid4().hex[:8]}.png"; fn.write_bytes(raw)
                path=f"/output/{fn.name}"; fmt="png"; data=path
        except Exception:
            pass  # 回退 SVG
    if not path:
        svg = _build_svg(prompt, title); fn = OUT/f"cover_{uuid.uuid4().hex[:8]}.svg"; fn.write_text(svg, encoding="utf-8")
        path=f"/output/{fn.name}"; fmt="svg"; data="data:image/svg+xml;utf8,"+urllib.parse.quote(svg)
    # 改动12：按会话记录产物，供对话界面缩略图条展示
    try: record_artifact(session_id, {"type":"image","title":title,"path":path,"thumb":path,
                                      "ts":int(datetime.datetime.now().timestamp())})
    except Exception: pass
    return {"format":fmt,"title":title,"path":path,"data":data}

# ---------- 终端 / 代码（真实执行，带安全加固） ----------
# 破坏性命令黑名单：覆盖 rm -rf ~ / ./ / * 等此前漏掉的写法
BLOCKLIST = ["rm -rf","rm -r /","rmdir /","mkfs","shutdown","reboot",":(){",":(){:|:&",
            "dd if=","> /dev/sd","format c","format /","wipefs","mv / ","chmod -R 777 /",
            "sudo rm","del /q","rd /s","shred ","> /dev/sda","echo > /dev"]
WRITE_CMDS = ["rm ","mv ","cp ","dd ","mkfs","shutdown","reboot","chmod","chown","touch ",
              "mkdir ","sudo ","> ",">> ","wipefs","shred ","mkfs."]
def is_dangerous(cmd):
    low=cmd.lower()
    for b in BLOCKLIST:
        if b in low: return b.strip()
    return None
def _readonly_blocked(cmd):
    low=cmd.lower().lstrip()
    return any(low.startswith(w.strip()) or f" {w.strip()}" in (" "+low) for w in WRITE_CMDS)
def run_terminal(cmd):
    cmd=cmd.strip()
    if not cmd: return "[空命令]"
    danger=is_dangerous(cmd)
    if danger:
        return f"[已拦截危险命令] 出于安全考虑，『{danger}』类操作被禁用。可改为只读命令如 ls / df -h / du -sh。"
    if config.get("sandbox",{}).get("readonly", False) and _readonly_blocked(cmd):
        return "[只读模式] 当前开启只读沙箱，写操作（rm/cp/mv/dd 等）已禁用。可在设置里关闭只读模式。"
    try:
        r=subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=25, cwd=str(APP_DIR))
        out=(r.stdout or "")+(r.stderr or "")
        return out[:4000] or (f"[exit {r.returncode}]" if r.returncode else "[无输出]")
    except subprocess.TimeoutExpired: return "[超时] 命令执行超过 25 秒已被终止。"
    except Exception as e: return f"[错误] {e}"
def run_code(code):
    fn=profile_dir()/f"run_{uuid.uuid4().hex[:6]}.py"; fn.write_text(code, encoding="utf-8")
    try:
        r=subprocess.run(["python3",str(fn)], capture_output=True, text=True, timeout=20)
        return ((r.stdout or "")+(r.stderr or ""))[:4000] or "[无输出]"
    except subprocess.TimeoutExpired: return "[超时] 代码运行超过 20 秒。"
    except Exception as e: return f"[错误] {e}"
    finally:
        try: fn.unlink()
        except Exception: pass

# ---------- 离谱玩法 ----------
def handle_fun(text, session_id):
    if "下载" in text or "steam" in text.lower():
        return ("🎮 **下载玩法（终端联动）**\n我来调用终端帮你拉取资源。请在对话里给出下载链接，例如：\n> 下载 https://example.com/game.zip\n\n或直接说「清理电脑」「连 b站」试试别的离谱玩法。")
    if "清理" in text or "垃圾" in text:
        out=run_terminal("echo '🧹 安全清理预览（仅统计，不删除）' && du -sh ~ 2>/dev/null; df -h / | tail -1")
        return f"🧹 **清理电脑（只读扫描）**\n```\n{out}\n```\n_默认只做统计，保护你的数据；真要删临时文件可指定路径。_"
    if "游戏" in text or "植物大战僵尸" in text or "pvz" in text.lower() or "贪吃蛇" in text or "搓" in text:
        html=gen_game("贪吃蛇"); fn=OUT/f"game_{uuid.uuid4().hex[:8]}.html"; fn.write_text(html, encoding="utf-8")
        # 改动12：记录游戏产物
        try: record_artifact(session_id, {"type":"game","title":"贪吃蛇（HTML 小游戏）","path":f"/output/{fn.name}","thumb":None,
                                          "ts":int(datetime.datetime.now().timestamp())})
        except Exception: pass
        return {"reply":"🕹️ **搓游戏玩法**\n已为你现场生成《贪吃蛇》可玩 HTML！点下面按钮在浏览器打开即玩：",
                "mode":"game","game_path":f"/output/{fn.name}"}
    if "b站" in text or "bilibili" in text.lower() or "连接" in text:
        out=run_terminal("curl -sI --max-time 10 https://www.bilibili.com | head -5")
        return f"🌐 **连接外部网站（终端 curl）**\n```\n{out}\n```"
    return None

# ---------- 联网搜索 ----------
def web_search(q, n=5):
    """联网搜索：返回 {"snippet": 摘要文本, "sources": [{title, url}]}。"""
    try:
        from bs4 import BeautifulSoup
        headers={"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        r=requests.post("https://html.duckduckgo.com/html/", data={"q":q},
                        headers=headers, timeout=15)
        soup=BeautifulSoup(r.text,"html.parser")
        items=[]; sources=[]
        for res in soup.select(".result")[:n]:
            a=res.select_one("a.result__a"); sn=res.select_one("a.result__snippet")
            if a:
                href=a.get("href","")
                real=href
                if "uddg=" in href:
                    try: real=parse_qs(urlparse(href).query).get("uddg",[href])[0]
                    except Exception: real=href
                title=a.get_text(strip=True)
                snippet=(sn.get_text(strip=True) if sn else "")
                items.append(f"- {title}：{snippet}")
                sources.append({"title":title, "url":real})
        if items:
            return {"snippet":"（联网搜索「"+q+"」）\n"+ "\n".join(items), "sources":sources}
        return {"snippet":f"[搜索无结果] 未能获取「{q}」的搜索结果，可能是网络受限。", "sources":[]}
    except Exception as e:
        return {"snippet":f"[搜索失败] {e} —— 当前环境可能无外网或搜索被拦截。", "sources":[]}

# ---------- 文档理解（RAG，本地召回） ----------
def load_docs():
    return _load(profile_dir()/"docs.json", [])
def save_docs(d):
    _save(profile_dir()/"docs.json", d)
def extract_text(fn, suffix):
    try:
        if suffix==".pdf":
            out=subprocess.run(["pdftotext",str(fn),"-"], capture_output=True, text=True, timeout=30)
            return out.stdout
        if suffix==".docx":
            from docx import Document as Docx
            return "\n".join(p.text for p in Docx(str(fn)).paragraphs)
        if suffix in (".txt",".md",".csv"):
            return fn.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return None
    return None
def add_doc(name, text):
    d=load_docs()
    chunks=re.split(r"(?<=[。！？\n])", text)
    chunks=[c.strip() for c in chunks if len(c.strip())>8]
    rec={"id":uuid.uuid4().hex[:8],"name":name,"text":text[:80000],"chunks":chunks[:600]}
    d.append(rec); save_docs(d); return rec
def retrieve_doc(q, k=3):
    d=load_docs()
    if not d: return ""
    # 中文字符级重叠打分（中文无空格，不能用分词）
    qchars=[c for c in q if c.strip()]
    if not qchars: return ""
    scored=[]
    for doc in d:
        for ch in doc["chunks"]:
            sc=sum(1 for c in qchars if c in ch)
            if sc>0: scored.append((sc, doc["name"], ch))
    scored.sort(reverse=True)
    if not scored: return ""
    return "（参考文档片段）\n"+"\n".join(f"[{name}] {ch[:280]}" for _,name,ch in scored[:k])

def rag_ctx(text):
    """把知识库检索结果注入上下文：用户上传过文档且命中相关片段时返回片段文本，否则返回空串。
    这样在普通聊天 / 写作里说『总结这份合同』也会真正用到文档，而不必走 Agent。"""
    try:
        rel = retrieve_doc(text, k=3)
    except Exception:
        rel = ""
    return rel

# ---------- 搓游戏（真实可玩 HTML） ----------
SNAKE_HTML = """<!doctype html><html lang="zh"><head><meta charset="utf-8">
<title>墨核 · 贪吃蛇</title><style>body{background:#0f2027;color:#fff;font-family:sans-serif;text-align:center}
canvas{background:#1b3a4b;border-radius:8px;margin-top:10px}#s{color:#7c5cff}</style></head>
<body><h2>🐍 贪吃蛇（墨核 AI 搓的）</h2><p id="s">方向键 / WASD 控制，吃到方块得分</p>
<canvas id="c" width="360" height="360"></canvas>
<script>const c=document.getElementById('c'),x=c.getContext('2d');let s=[[180,180]],d=[20,0],f=[100,100],sc=0;
function loop(){let h=[s[0][0]+d[0],s[0][1]+d[1]];
 if(h[0]<0||h[0]>340||h[1]<0||h[1]>340||s.some(p=>p[0]==h[0]&&p[1]==h[1])){alert('游戏结束，得分 '+sc);s=[[180,180]];d=[20,0];sc=0;return;}
 s.unshift(h);if(h[0]==f[0]&&h[1]==f[1]){sc++;f=[Math.floor(Math.random()*17)*20,Math.floor(Math.random()*17)*20];}else s.pop();
 x.clearRect(0,0,360,360);x.fillStyle='#7c5cff';s.forEach(p=>x.fillRect(p[0],p[1],18,18));x.fillStyle='#ffd166';x.fillRect(f[0],f[1],18,18);
 document.getElementById('s').textContent='得分：'+sc;setTimeout(loop,120);}
 addEventListener('keydown',e=>{if(e.key=='ArrowUp')d=[0,-20];if(e.key=='ArrowDown')d=[0,20];if(e.key=='ArrowLeft')d=[-20,0];if(e.key=='ArrowRight')d=[20,0];});
 loop();</script></body></html>"""
def gen_game(kind="贪吃蛇"):
    return SNAKE_HTML

# ---------- 多模态视觉（需配置 vision） ----------
def vision_understand(prompt, image_b64):
    c=config.get("vision") or {}
    if not (c.get("enabled") and c.get("api_key")):
        return None
    return call_api("vision", "请用中文描述这张图片，并结合用户问题作答。", None, prompt, image_b64)

# ---------- Agent 工具集（本地真实编排，LLM 增强） ----------
AGENT_TOOLS = {
    "search":   {"desc":"联网搜索","run":lambda q: web_search(q)},
    "image":    {"desc":"AI 生图","run":lambda q: gen_image(q,"agent")},
    "code":     {"desc":"运行 Python","run":lambda q: run_code(q)},
    "terminal": {"desc":"执行命令","run":lambda q: run_terminal(q)},
    "doc":      {"desc":"基于文档","run":lambda q: retrieve_doc(q)},
    "write":    {"desc":"写作","run":lambda q: gen_writing(q,"agent")},
}
def agent_step(step_text):
    t=step_text
    if any(k in t for k in ["搜索","查一下","查 ","搜 ","资料","最新","资讯"]): return ("search", AGENT_TOOLS["search"]["run"](t))
    if any(k in t for k in ["生图","封面","画","图片"]): return ("image", AGENT_TOOLS["image"]["run"](t))
    if any(k in t for k in ["代码","python","运行","脚本"]): return ("code", AGENT_TOOLS["code"]["run"](t))
    if any(k in t for k in ["命令","终端","执行","ls","df"]): return ("terminal", AGENT_TOOLS["terminal"]["run"](t))
    if any(k in t for k in ["文档","资料","基于","总结这篇","读"]): return ("doc", AGENT_TOOLS["doc"]["run"](t))
    return ("write", AGENT_TOOLS["write"]["run"](t))  # 默认写作
def llm_pick_tool(text):
    """LLM 开启时，让模型为某一步骤挑最合适的工具（function-calling 思路的轻量实现）。"""
    if not (config["llm"].get("enabled") and config["llm"].get("api_key")):
        return None
    prompt = ('你是 Agent 工具选择器。仅回复一个 JSON：{"tool":"search|image|code|terminal|doc|write"}，'
              '根据用户步骤描述选择最合适的一个工具。不要输出其他内容。')
    raw = call_llm(prompt, None, "步骤："+text)
    if not raw or raw.startswith("[LLM"): return None
    m = re.search(r"\{.*\}", raw, re.DOTALL)
    if not m: return None
    try:
        d = json.loads(m.group(0)); t = d.get("tool")
        if t in AGENT_TOOLS: return t
    except Exception: pass
    return None
def agent_run(goal, steps):
    use_llm = config["llm"].get("enabled") and config["llm"].get("api_key")
    log=[f"[Agent] 目标：{goal}"]
    collected=[]
    for i,s in enumerate(steps):
        tool = llm_pick_tool(s) if use_llm else None
        if not tool:
            tool,_ = agent_step(s)            # 关键词兜底选工具
        out = AGENT_TOOLS[tool]["run"](s)     # 真实调用工具
        if isinstance(out, dict): out = out.get("snippet","")
        collected.append((tool, s, str(out)))
        log.append(f"→ 步骤{i+1}（工具：{tool}）{s}\n{str(out)[:500]}")
    # LLM 开启时，把各步骤产出合成为一份最终交付物
    if use_llm and collected:
        synth_in = "\n\n".join(f"[{t}] 步骤：{q}\n{out[:400]}" for t,q,out in collected)
        r = call_api("llm", SYSTEM_BASE + "你是 Agent 总结器：请基于下面各步骤的工具产出，整合成一份条理清晰、可直接使用的最终交付物（结论/文稿/报告）。只输出交付物本身。",
                     None, "目标："+goal+"\n\n"+synth_in)
        if r and not r.startswith("[LLM"):
            log.append("\n— — —\n[Agent] 智能合成结果：\n"+r)
    log.append("\n[Agent] 完成。")
    return "\n".join(log)

# ---------- 意图路由 ----------
WELCOME = ("我是**墨核 AI** 🐱，你的桌面智能助手：\n"
           "- **自由聊天 / 问答 / 搜资料**（联网搜索真实信息）\n- **写作**：小说 / 剧本 / 文章 / 文案（这是我的强项）\n"
           "- **AI 生图 / 生图提示词**、**写代码**、**跑终端命令**\n- 载入 **Skill** 插件、调用 **记忆**、当 **Agent** 帮你做项目\n\n"
           "试试：「你都能帮我做什么」「搜索今天的热点」「写一段仙侠短篇，主角叫苏无尘」「画一张赛博朋克封面」\n"
           "（未配置大模型时我用本地模板应答；在右上角「⚙️ 设置」填入 API 即可解锁完整对话与搜索能力。）")

# ---------- 语义路由（LLM 开启时） ----------
def llm_classify(text, history, session_id=None):
    """调用 LLM 做意图分类，返回 {'intent':..., 'need_clarify':bool} 或 None。"""
    cfg = get_llm_cfg(session_id)
    if not (_llm_usable(cfg)):
        return None
    raw = call_llm(SYSTEM_CLASSIFY, history, "用户："+text, cfg=cfg)
    if not raw or raw.startswith("[LLM"): return None
    m = re.search(r"\{.*\}", raw, re.DOTALL)
    if not m: return None
    try:
        d = json.loads(m.group(0))
        if isinstance(d, dict) and "intent" in d:
            d.setdefault("need_clarify", False); return d
    except Exception: pass
    return None

def clarify_prompt(text):
    return ("🐱 我有点没 get 到你的意思～ 你大概是想做哪类事？直接选一个说就行：\n"
            "- **写点东西**：小说 / 公众号文章 / 文案 / 剧本\n"
            "- **出图**：AI 生图 / 生图提示词\n"
            "- **查资料**：联网搜索 / 读本地文档\n"
            "- **动手**：跑终端命令 / 写 Python\n"
            "- **玩**：搓游戏 / 角色扮演 / 连外部站\n"
            "- **其他**：翻译 / 视觉理解 / 让我当 Agent 帮你做项目\n\n"
            f"你刚说的「{text[:30]}」我记下了，补一句方向我就开工。")

def route_by_intent(intent, text, session_id, history):
    """把 LLM 分类出的意图映射到对应的处理分支（仅在关键词路由都未命中时调用）。"""
    intent = (intent or "").lower()
    cfg = get_llm_cfg(session_id)
    if intent == "writing":
        return _wrap_stream(llm_or(SYSTEM_WRITING, memory_context(session_id)+"用户需求："+text+"\n"+rag_ctx(text),
                                lambda: gen_writing(text, session_id), session_id, history, stream=True), "writing")
    if intent == "translate":
        if _llm_usable(cfg):
            return {"reply_stream": call_api("llm", SYSTEM_TRANSLATE, None, text, cfg=cfg, stream=True), "mode":"translate"}
        return {"reply":"[翻译] 请先在设置里启用大模型（翻译需要 LLM）。原文：\n"+text, "mode":"translate"}
    if intent == "search":
        sr = web_search(text)
        if _llm_usable(cfg):
            return {"reply_stream": call_api("llm", SYSTEM_SEARCH, history, memory_context(session_id)+text+"\n"+sr["snippet"], cfg=cfg, stream=True), "mode":"search", "sources":sr.get("sources",[])}
        return {"reply":sr["snippet"], "mode":"search", "sources":sr.get("sources",[])}
    if intent == "image":
        return {"reply":"已生成封面：", "image": gen_image(text, session_id), "mode":"image"}
    if intent == "prompt":
        return _wrap_stream(llm_or(SYSTEM_PROMPT, memory_context(session_id)+text,
                                lambda: gen_prompt(text), session_id, history, stream=True), "prompt")
    if intent == "memory":
        return {"reply": format_memory(session_id), "mode":"memory"}
    if intent == "code":
        code = gen_code(text)
        if code:
            return {"reply":"已为你生成 Python 代码，并在右侧代码面板（已铺满）展开，点「运行」即可执行。",
                    "mode":"code_run", "code":code}
        return {"reply":"已为你打开代码面板，点「运行」即可执行 Python；也可 `/code print('hi')` 直接跑。", "mode":"code_hint"}
    if intent == "terminal":
        cmd = gen_terminal_cmd(text)
        if cmd:
            return {"reply":"已为你打开终端面板并铺满，命令已填入，回车即可执行（写操作会要求确认）：\n```bash\n"+cmd+"\n```",
                    "mode":"terminal_run", "command":cmd}
        return {"reply":"已为你打开终端面板，也可以在对话里用 `/terminal <命令>` 直接执行。", "mode":"terminal_hint"}
    if intent == "agent":
        plan = gen_agent_plan(text)
        if plan:
            steps_txt = "\n".join(f"- {s}" for s in plan.get("steps",[]))
            return {"reply":"已为你打开 Agent 面板并铺满，目标与步骤已生成：\n**目标**："+plan.get("goal",text)+"\n"+steps_txt,
                    "mode":"agent_run", "goal":plan.get("goal",text), "steps":plan.get("steps",[])}
        return {"reply":"已为你打开 Agent 面板，在右侧定义目标与多步骤，我会真实串联 搜索/生图/代码/终端/文档/写作 工具执行。", "mode":"agent_hint"}
    if intent == "vision":
        return {"reply":"点输入框左侧 📷 发一张图片，我就能看图理解（需先在设置里启用视觉模型）。", "mode":"vision_hint"}
    if intent == "fun":
        fun = handle_fun(text, session_id)
        return fun if fun else {"reply":"想玩点啥？可以说：搓一个贪吃蛇 / 清理电脑 / 连 b站 / 打开终端跑命令。", "mode":"fun"}
    if intent == "skill":
        return {"reply":"你想用哪个技能？当前内置：起名器 / 生图提示词。也可以在右侧「技能」面板自己创建提示词型或人设型技能。", "mode":"chat"}
    # chat / 其他 → 通用对话
    if _llm_usable(cfg):
        return {"reply_stream": call_llm(SYSTEM_BASE, history, rag_ctx(text)+"\n"+text, cfg=cfg, stream=True), "mode":"chat"}
    return None

def apply_persona(skill, session_id):
    personas[session_id]={"name":skill["name"],"content":skill["content"]}
    _save(profile_dir()/"personas.json", personas)
    return f"🐱 已切换人设：**{skill['name']}**。\n{skill['content'][:60]}……\n（直接跟我说话即可，输入「退出角色扮演」结束）"
def gen_from_skill(skill, text, session_id):
    if skill["id"]=="skill_naming":
        kinds=["苏","沈","陆","顾","叶","云","夜","白"]
        return skill["content"]+"\n\n示例输出：\n"+"\n".join([f"- {random.choice(kinds)}{random.choice(['无尘','惊鸿','长庚','挽星','听澜','昭明'])}：清冷出尘，适合隐世高人" for _ in range(4)])
    if skill["id"]=="skill_plot":
        return skill["content"]+"\n\nA 稳妥流：主角先守后攻；B 逆袭流：扮猪吃虎；C 黑马流：第三方势力入场搅局。"
    if skill["id"]=="skill_prompt_img": return gen_prompt(text)
    return skill["content"]
def format_memory(session_id):
    m=get_mem(session_id); lines=["# 🧠 记忆库（长篇创作防丢稿/防幻觉）",""]
    lines.append(f"- **风格偏好**：{m['style'] or '（未设定）'}")
    lines.append(f"- **人物**：{', '.join(m['characters']) or '（暂无）'}")
    lines.append(f"- **设定/事实**：{(chr(10)+'  - ').join(m['facts']) or '（暂无）'}")
    lines.append(f"- **世界观**：{', '.join(m['world']) or '（暂无）'}")
    if not any([m['style'],m['characters'],m['facts'],m['world']]):
        lines.append("\n_还在空白状态～ 告诉我主角名字、题材或风格，我会自动记住。_")
    return "\n".join(lines)

def route(text, session_id, history=None):
    t=text.strip()
    mem=get_mem(session_id)
    cfg=get_llm_cfg(session_id)
    if session_id in personas and ("退出" in t and "角色" in t):
        personas.pop(session_id); _save(profile_dir()/"personas.json", personas)
        return {"reply":"已退出角色扮演，回到全能助手模式喵～（不是）","mode":"chat"}
    if session_id in personas:
        p=personas[session_id]
        # 角色模式下仍优先识别明确的 搜索/生图/命令/代码 意图，做到"能聊天也能搜/画"
        if any(k in t for k in ["搜索","搜一下","查一下","查资料","最新资讯","搜 "]):
            sr=web_search(t)
            if _llm_usable(cfg):
                return {"reply_stream": call_api("llm", SYSTEM_SEARCH, history, (p["content"]+"\n")+memory_context(session_id)+t+"\n"+sr["snippet"], cfg=cfg, stream=True), "mode":"search","sources":sr.get("sources",[])}
            return {"reply":sr["snippet"],"mode":"search","sources":sr.get("sources",[])}
        if any(k in t for k in ["画","生图","封面","做图","出图","配图","图片","图一张","生成图"]):
            return {"reply":"已生成封面：","image":gen_image(t,session_id),"mode":"image"}
        if t.startswith("/terminal") or t.startswith("/term"):
            out=run_terminal(t.split(" ",1)[1] if " " in t else ""); return {"reply":f"```bash\n{out}\n```","mode":"terminal"}
        if t.startswith("/code"):
            code=t.split(" ",1)[1] if " " in t else "print('hello')"; return {"reply":f"```python\n{run_code(code)}\n```","mode":"code"}
        if _llm_usable(cfg):
            return {"reply_stream": call_llm(p["content"], history, t, cfg=cfg, stream=True), "mode":"roleplay"}
        return {"reply":f"（{p['name']}）{t}……嗯哼，{random.choice(['人家有点想你了喵～','才、才不是担心你呢','要喝杯茶吗？'])}","mode":"roleplay"}
    if t.startswith("/image"):
        return {"reply":"已生成封面：","image":gen_image(t[6:],session_id),"mode":"image"}
    if t.startswith("/terminal") or t.startswith("/term"):
        out=run_terminal(t.split(" ",1)[1] if " " in t else ""); return {"reply":f"```bash\n{out}\n```","mode":"terminal"}
    if t.startswith("/code"):
        code=t.split(" ",1)[1] if " " in t else "print('hello')"; return {"reply":f"```python\n{run_code(code)}\n```","mode":"code"}
    if t.startswith("/skill"):
        return {"reply":"打开右侧『技能』面板即可创建/管理 Skill；或在对话里直接触发某个技能关键词。","mode":"chat"}
    for s in skills:
        if any(k.lower() in t.lower() for k in s.get("trigger",[])):
            if s.get("type")=="persona":
                return {"reply":apply_persona(s,session_id),"mode":"roleplay"}
            if _llm_usable(cfg):
                g = call_llm(s["content"], history, memory_context(session_id)+"用户需求："+t, cfg=cfg, stream=True)
                def _skill_wrap(g, name):
                    yield f"⚡ 已调用技能 **{name}**\n\n"
                    yield from g
                return {"reply_stream": _skill_wrap(g, s["name"]), "mode":"skill"}
            return {"reply":f"⚡ 已调用技能 **{s['name']}**\n\n{gen_from_skill(s,t,session_id)}","mode":"skill"}
    fun=handle_fun(t,session_id)
    if fun is not None:
        if isinstance(fun, dict): return fun
        return {"reply":fun,"mode":"fun"}
    if any(k in t for k in ["搜索","搜一下","查一下","查资料","最新资讯","搜 "]):
        sr=web_search(t)
        if _llm_usable(cfg):
            return {"reply_stream": call_api("llm", SYSTEM_SEARCH, history, memory_context(session_id)+t+"\n"+sr["snippet"], cfg=cfg, stream=True), "mode":"search","sources":sr.get("sources",[])}
        return {"reply":sr["snippet"],"mode":"search","sources":sr.get("sources",[])}
    if any(k in t for k in ["翻译","译一下","中英","英译中","中译英"]):
        if _llm_usable(cfg):
            return {"reply_stream": call_api("llm", SYSTEM_TRANSLATE, None, t, cfg=cfg, stream=True), "mode":"translate"}
        return {"reply":"[翻译] 请先在设置里启用大模型（翻译需要 LLM）。原文：\n"+t,"mode":"translate"}
    if any(k in t for k in ["生图","封面","插图","画","图片"]):
        return {"reply":"已生成封面：","image":gen_image(t,session_id),"mode":"image"}
    if any(k in t for k in ["提示词","生图词","prompt"]):
        return _wrap_stream(llm_or(SYSTEM_PROMPT, memory_context(session_id)+t, lambda:gen_prompt(t), session_id, history, stream=True),"prompt")
    if "终端" in t or "命令" in t or "执行" in t:
        return {"reply":"已为你打开终端面板，也可以在对话里用 `/terminal <命令>` 直接执行，例如 `/terminal ls -la`。","mode":"terminal_hint"}
    if any(k in t for k in ["代码","python","运行"]):
        return {"reply":"已为你打开代码面板，点「运行」即可执行 Python；也可 `/code print('hi')` 直接跑。","mode":"code_hint"}
    if "记忆" in t:
        return {"reply":format_memory(session_id),"mode":"memory"}
    if detect_genre(t) or any(k in t for k in ["写","创作","生成","来一篇","写个","帮我"]):
        return _wrap_stream(llm_or(SYSTEM_WRITING, memory_context(session_id)+"用户需求："+t+"\n"+rag_ctx(t), lambda:gen_writing(t,session_id), session_id, history, stream=True),"writing")
    # 兜底：LLM 开启 → 语义分类路由；含糊输入 → 澄清反问，而非直接丢欢迎语
    if _llm_usable(cfg):
        cls = llm_classify(t, history, session_id)
        if cls and cls.get("need_clarify"):
            return {"reply": clarify_prompt(t), "mode":"clarify"}
        if cls and cls.get("intent"):
            routed = route_by_intent(cls["intent"], t, session_id, history)
            if routed: return routed
        return {"reply_stream": call_llm(SYSTEM_BASE, history, rag_ctx(t)+"\n"+t, cfg=cfg, stream=True), "mode":"chat"}
    return {"reply":WELCOME,"mode":"chat"}

# ---------- Word 导出 ----------
def md_to_docx(doc, text):
    fence=False
    for line in text.split("\n"):
        line=line.rstrip()
        if line.strip()=="":
            continue
        if line.startswith("```"):
            fence=not fence; continue
        if fence:
            p=doc.add_paragraph(); r=p.add_run(line); r.font.name="Consolas"; continue
        if line.startswith("### "): doc.add_heading(line[4:],level=3); continue
        if line.startswith("## "): doc.add_heading(line[3:],level=2); continue
        if line.startswith("# "): doc.add_heading(line[2:],level=1); continue
        if line.startswith("> "):
            p=doc.add_paragraph(); r=p.add_run(line[2:]); r.italic=True; continue
        if line.startswith("- ") or line.startswith("* "):
            p=doc.add_paragraph(style="List Bullet"); _add_runs(p,line[2:]); continue
        if line.startswith("|"):
            p=doc.add_paragraph(); _add_runs(p,line); continue
        p=doc.add_paragraph(); _add_runs(p,line)
def _add_runs(p, text):
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            r=p.add_run(part[2:-2]); r.bold=True
        else:
            p.add_run(part)

# ---------- Flask ----------
app = Flask(__name__, static_folder=str(STATIC), template_folder=str(STATIC))
@app.route("/")
def index(): return send_from_directory(str(STATIC),"index.html")
@app.route("/output/<path:p>")
def output_file(p): return send_from_directory(str(OUT), p)
@app.route("/api/chat", methods=["POST"])
def api_chat():
    data=request.json or {}; text=(data.get("message") or "").strip()
    session_id=data.get("session_id") or str(uuid.uuid4())
    if not text: return jsonify({"error":"empty"}),400
    conv=conversations.setdefault(session_id, {"title":text[:20],"messages":[],"group":(data.get("group") or "")})
    history=list(conv.get("messages",[]))
    update_memory(session_id, text)
    res=route(text, session_id, history)
    # 兼容：route() 在 LLM 文本分支会返回 reply_stream 生成器，这里一次性消费成完整文本
    if "reply_stream" in res:
        try:
            full = "".join(res["reply_stream"])
        except Exception as e:
            full = f"[LLM 调用失败] {e}"
        res = dict(res); res["reply"] = full; res.pop("reply_stream", None)
    conv["messages"].append({"role":"user","content":text})
    conv["messages"].append({"role":"assistant","content":res.get("reply","")})
    conversations[session_id]=conv
    _save(profile_dir()/"conversations.json", conversations)
    extra={k:v for k,v in res.items() if k not in ("reply","image","mode")}
    return jsonify({"session_id":session_id,"reply":res.get("reply",""),"image":res.get("image"),
                    "mode":res.get("mode","chat"),"memory":get_mem(session_id),**extra})

@app.route("/api/chat_stream", methods=["POST"])
def api_chat_stream():
    """SSE 流式对话：逐块推送 token，支持分支（编辑/重新生成）与中断（客户端断开即停止）。"""
    data=request.json or {}; text=(data.get("message") or "").strip()
    session_id=data.get("session_id") or str(uuid.uuid4())
    if not text: return jsonify({"error":"empty"}),400
    keep = data.get("keep")
    regen = bool(data.get("regen"))
    append_user = True
    if keep is not None:
        try: keep=int(keep)
        except Exception: keep=None
    if keep is not None:
        conv = conversations.get(session_id)
        if not conv: return jsonify({"error":"no conversation"}),404
        msgs = list(conv.get("messages", []))
        if regen:
            # 重新生成：保留 keep 之前（含 keep-1 的用户消息），丢弃 keep 处的助手消息及其后
            if keep <= 0 or keep > len(msgs): return jsonify({"error":"bad index"}),400
            history = msgs[:keep-1]
            conv["messages"] = msgs[:keep]
            append_user = False
        else:
            # 编辑：丢弃 keep 处的用户消息及其后，用新文本作为用户消息重新生成
            history = msgs[:keep]
            conv["messages"] = msgs[:keep]
    else:
        conv = conversations.setdefault(session_id, {"title":text[:20],"messages":[],"group":(data.get("group") or "")})
        history = list(conv.get("messages", []))
    update_memory(session_id, text)
    res = route(text, session_id, history)
    image = res.get("image")
    game_path = res.get("game_path")
    sources = res.get("sources", [])
    mode = res.get("mode", "chat")
    meta = {"type":"meta","session_id":session_id,"mode":mode,
            "image":image,"game_path":game_path,"sources":sources}
    def event(d): return "data: "+json.dumps(d, ensure_ascii=False)+"\n\n"
    def gen_events():
        acc = ""
        try:
            yield event(meta)
            if "reply_stream" in res:
                for chunk in res["reply_stream"]:
                    acc += chunk
                    yield event({"type":"token","text":chunk})
            else:
                reply = res.get("reply", "")
                acc = reply
                yield event({"type":"token","text":reply})
            yield event({"type":"done","text":acc})
        finally:
            # 无论正常结束还是客户端中断（GeneratorExit），都尽量保存已生成内容
            try:
                if append_user:
                    conv["messages"].append({"role":"user","content":text})
                conv["messages"].append({"role":"assistant","content":acc})
                conversations[session_id]=conv
                _save(profile_dir()/"conversations.json", conversations)
            except Exception:
                pass
    return Response(stream_with_context(gen_events()), mimetype="text/event-stream",
                    headers={"Cache-Control":"no-cache","X-Accel-Buffering":"no"})
@app.route("/api/terminal", methods=["POST"])
def api_terminal():
    return jsonify({"output":run_terminal((request.json or {}).get("command",""))})
@app.route("/api/code", methods=["POST"])
def api_code():
    return jsonify({"output":run_code((request.json or {}).get("code",""))})
@app.route("/api/image", methods=["POST"])
def api_image():
    d=request.json or {}; return jsonify(gen_image(d.get("prompt",""), d.get("session_id","default")))
@app.route("/api/skills", methods=["GET","POST","DELETE"])
def api_skills():
    global skills
    if request.method=="GET": return jsonify(skills)
    if request.method=="POST":
        s=request.json; s["id"]=s.get("id") or f"skill_{uuid.uuid4().hex[:8]}"; s.setdefault("icon","🔧")
        skills.append(s); _save(profile_dir()/"skills.json", skills); return jsonify(s)
    sid=request.args.get("id"); skills=[x for x in skills if x["id"]!=sid]; _save(profile_dir()/"skills.json", skills)
    return jsonify({"ok":True})
@app.route("/api/memory")
def api_memory():
    return jsonify(get_mem(request.args.get("session_id","default")))
@app.route("/api/artifacts")
def api_artifacts():
    """返回某会话的全部产物（按会话分目录记录），供对话界面缩略图条展示。"""
    sid = request.args.get("session_id") or "default"
    return jsonify(load_artifacts().get(sid, []))
@app.route("/api/conversations", methods=["GET","POST","DELETE"])
def api_conversations():
    if request.method=="DELETE":
        sid=request.args.get("id")
        if sid in conversations:
            del conversations[sid]; _save(profile_dir()/"conversations.json", conversations)
        return jsonify({"ok":True})
    if request.method=="POST":
        d=request.json or {}; sid=d.get("id"); title=(d.get("title") or "").strip()
        if sid in conversations:
            if title: conversations[sid]["title"]=title
            if "group" in d: conversations[sid]["group"]=d["group"]
            _save(profile_dir()/"conversations.json", conversations)
        return jsonify({"ok":True})
    return jsonify([{"id":k,"title":v.get("title",""),"count":len(v.get("messages",[])),"model":v.get("model"),"group":v.get("group","")} for k,v in conversations.items()])

@app.route("/api/conversations/<sid>")
def api_conversation_detail(sid):
    """返回单个对话的完整消息，供前端回看历史。"""
    conv = conversations.get(sid)
    if not conv:
        return jsonify({"error":"not found"}), 404
    return jsonify({"id":sid, "title":conv.get("title",""), "model":conv.get("model"),
                    "messages":conv.get("messages",[])})

@app.route("/api/search")
def api_search():
    """对全部对话的消息内容做关键词检索，返回匹配对话与命中的片段。"""
    q = (request.args.get("q") or "").strip()
    if not q:
        return jsonify([])
    res = []
    for k, v in conversations.items():
        hit = None
        for m in v.get("messages", []):
            c = m.get("content", "")
            idx = c.find(q)
            if idx >= 0:
                s = max(0, idx-20); hit = c[s:idx+len(q)+40].replace("\n", " ")
                break
        if hit or q in (v.get("title","") or ""):
            res.append({"id":k, "title":v.get("title",""), "snippet":hit or v.get("title","")})
    return jsonify(res)
@app.route("/api/conversations/model", methods=["POST"])
def api_conv_model():
    d=request.json or {}; sid=d.get("session_id")
    if not sid: return jsonify({"error":"no session"}),400
    model=(d.get("model") or "").strip()
    conv=conversations.setdefault(sid, {"title":"对话","messages":[]})
    if model:
        conv["model"]=model
    else:
        conv.pop("model", None)
    _save(profile_dir()/"conversations.json", conversations)
    return jsonify({"ok":True,"model":conv.get("model")})
@app.route("/api/config", methods=["GET","POST"])
def api_config():
    global config
    config.setdefault("provider_keys", {})
    if request.method=="POST":
        patch=request.json or {}
        # 选择内置供应商时，自动套用其 base_url / model 预设；前端传了 model 则优先用
        if patch.get("provider") and patch["provider"] in PROVIDERS:
            p=PROVIDERS[patch["provider"]]
            config["provider"]=patch["provider"]
            chosen_model = patch.get("model") or p["model"]
            config["llm"]["base_url"]=p["base_url"]; config["llm"]["model"]=chosen_model
            config["vision"]["base_url"]=p["base_url"]; config["vision"]["model"]=chosen_model
            # 若该供应商此前记住过自己的 Key，则恢复（切换不丢失）
            pk = config["provider_keys"].get(config["provider"])
            if isinstance(pk, dict):
                if pk.get("api_key"): config["llm"]["api_key"]=pk["api_key"]
                if pk.get("base_url"): config["llm"]["base_url"]=pk["base_url"]
                if pk.get("model"): config["llm"]["model"]=pk["model"]
        for sec in ("llm","image","vision"):
            if sec in patch and isinstance(patch[sec], dict): config[sec].update(patch[sec])
        # #7 切换/保存时，把 LLM 的 Key+Base 继承到生图与视觉（用户勾选时）
        if patch.get("inherit_keys") and config["llm"].get("api_key"):
            for sec in ("image","vision"):
                config[sec]["api_key"]=config["llm"]["api_key"]
                if config["llm"].get("base_url"): config[sec]["base_url"]=config["llm"]["base_url"]
        # #5 把当前供应商的 Key 单独存一份，切换回来不丢
        if "llm" in patch and isinstance(patch["llm"], dict) and "api_key" in patch["llm"]:
            prov = config.get("provider","custom")
            config["provider_keys"].setdefault(prov, {})
            config["provider_keys"][prov]["api_key"] = patch["llm"]["api_key"]
            if config["llm"].get("base_url"): config["provider_keys"][prov]["base_url"] = config["llm"]["base_url"]
            if config["llm"].get("model"): config["provider_keys"][prov]["model"] = config["llm"]["model"]
        if "sandbox" in patch and isinstance(patch["sandbox"], dict): config["sandbox"].update(patch["sandbox"])
        if "provider_keys" in patch and isinstance(patch["provider_keys"], dict): config["provider_keys"].update(patch["provider_keys"])
        if "active_profile" in patch: config["active_profile"]=patch["active_profile"]
        if "hotkey" in patch and isinstance(patch["hotkey"], str): config["hotkey"]=patch["hotkey"].strip()
        save_config()
    return jsonify({"provider":config.get("provider","custom"),"providers":PROVIDERS,
                    "llm":config["llm"],"image":config["image"],"vision":config["vision"],
                    "provider_keys":config.get("provider_keys",{}),
                    "sandbox":config["sandbox"],"active_profile":config.get("active_profile","default"),
                    "hotkey":config.get("hotkey","ctrl+alt+m"),
                    "version":APP_VERSION,"update_url":UPDATE_URL_DEFAULT})
@app.route("/api/ollama/check")
def api_ollama_check():
    """探测本机是否运行了 Ollama（零 Key 本地推理的前提）。"""
    try:
        r = requests.get("http://localhost:11434/api/tags", timeout=2)
        if r.ok:
            tags = r.json().get("models", [])
            return jsonify({"running": True, "models": [m.get("name") for m in tags][:10]})
        return jsonify({"running": False, "models": []})
    except Exception:
        return jsonify({"running": False, "models": []})

_MODEL_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9:._\-]{0,63}$")
def _ollama_cli(sub, *args, timeout=60):
    if not _MODEL_RE.match(sub) and sub not in ("pull","rm","list","serve"):
        return None, "非法参数"
    try:
        r = subprocess.run(["ollama", sub, *args], capture_output=True, text=True, timeout=timeout)
        return (r.stdout or "")+(r.stderr or ""), None
    except subprocess.TimeoutExpired:
        return None, f"命令超时（>{timeout}s）"
    except Exception as e:
        return None, str(e)

@app.route("/api/ollama/models")
def api_ollama_models():
    """列出本机已安装的 Ollama 模型。"""
    try:
        r = requests.get("http://localhost:11434/api/tags", timeout=3)
        if r.ok:
            return jsonify({"running": True, "models": [m.get("name") for m in r.json().get("models", [])]})
    except Exception:
        pass
    return jsonify({"running": False, "models": []})

@app.route("/api/ollama/pull", methods=["POST"])
def api_ollama_pull():
    d = request.json or {}
    name = (d.get("model") or "").strip()
    if not _MODEL_RE.match(name):
        return jsonify({"ok": False, "error": "模型名非法（仅限字母数字 : . _ -）"}), 400
    out, err = _ollama_cli("pull", name, timeout=900)
    if err:
        return jsonify({"ok": False, "error": err})
    return jsonify({"ok": True, "output": (out or "")[-3000:]})

@app.route("/api/ollama/delete", methods=["POST"])
def api_ollama_delete():
    d = request.json or {}
    name = (d.get("model") or "").strip()
    if not _MODEL_RE.match(name):
        return jsonify({"ok": False, "error": "模型名非法"}), 400
    out, err = _ollama_cli("rm", name, timeout=120)
    if err:
        return jsonify({"ok": False, "error": err})
    return jsonify({"ok": True, "output": (out or "")[:1500]})

@app.route("/api/ollama/use", methods=["POST"])
def api_ollama_use():
    d = request.json or {}
    name = (d.get("model") or "").strip()
    if not name:
        return jsonify({"error": "缺少模型名"}), 400
    PROV = PROVIDERS.get("ollama", {})
    config["provider"] = "ollama"
    config["llm"]["enabled"] = True
    config["llm"]["base_url"] = PROV.get("base_url", "http://localhost:11434/v1")
    config["llm"]["model"] = name
    config["llm"]["api_key"] = ""          # 本地推理无需 Key
    save_config()
    return jsonify({"ok": True, "model": name})

@app.route("/api/profiles", methods=["GET","POST"])
def api_profiles():
    base=DATA_HOME/"profiles"
    if request.method=="POST":
        name=(request.json or {}).get("name","").strip() or "default"
        profile_dir(name); config["active_profile"]=name; save_config(); load_profile()
        return jsonify({"active":name,"profiles":sorted(p.name for p in base.iterdir() if p.is_dir())})
    return jsonify({"active":config.get("active_profile","default"),
                    "profiles":sorted(p.name for p in base.iterdir() if p.is_dir())})
@app.route("/api/export/docx", methods=["POST"])
def api_export_docx():
    sid=(request.json or {}).get("session_id")
    conv=conversations.get(sid)
    if not conv: return jsonify({"error":"no conversation"}),404
    doc=Document(); doc.add_heading(conv.get("title","墨核 AI 创作") or "墨核 AI 创作", level=0)
    for msg in conv.get("messages",[]):
        role="你" if msg["role"]=="user" else "墨核 AI"
        h=doc.add_paragraph(); rr=h.add_run(f"【{role}】"); rr.bold=True
        md_to_docx(doc, msg["content"])
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    name=(re.sub(r"[^\w一-龥-]","",conv.get("title","墨核AI创作")) or "墨核AI创作")+".docx"
    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name=name)

# ---------- 文档上传 / 检索（RAG） ----------
@app.route("/api/docs", methods=["GET","POST","DELETE"])
def api_docs():
    if request.method=="GET":
        return jsonify([{"id":d["id"],"name":d["name"],"chunks":len(d["chunks"])} for d in load_docs()])
    if request.method=="DELETE":
        did=request.args.get("id"); save_docs([x for x in load_docs() if x["id"]!=did])
        return jsonify({"ok":True})
    f=request.files.get("file")
    if not f: return jsonify({"error":"no file"}),400
    suffix=Path(f.filename).suffix.lower()
    tmp=profile_dir()/f"up_{uuid.uuid4().hex[:8]}{suffix}"; f.save(str(tmp))
    text=extract_text(tmp, suffix)
    try: tmp.unlink()
    except Exception: pass
    if not text or not text.strip():
        return jsonify({"error":"无法提取文本（仅支持 PDF/Word/TXT/MD/CSV）"}),415
    rec=add_doc(f.filename, text)
    return jsonify({"ok":True,"id":rec["id"],"name":rec["name"],"chunks":len(rec["chunks"])})

# ---------- 多模态视觉 ----------
@app.route("/api/vision", methods=["POST"])
def api_vision():
    data=request.json or {}
    prompt=data.get("prompt","描述这张图片")
    img=data.get("image","")
    if img.startswith("data:"): img=img.split(",",1)[1]
    res=vision_understand(prompt, img)
    if res is None:
        return jsonify({"reply":"[视觉] 请先在设置里启用「视觉模型」（多模态 LLM）并填入 API Key。","mode":"vision"})
    return jsonify({"reply":res,"mode":"vision"})

# ---------- 导出：Markdown / PDF / 公众号 ----------
def _conv_md(conv):
    lines=[f"# {conv.get('title','墨核 AI 创作')}",""]
    for m in conv.get("messages",[]):
        role="你" if m["role"]=="user" else "墨核 AI"
        lines.append(f"**{role}：**\n{m['content']}\n")
    return "\n".join(lines)
@app.route("/api/export/md", methods=["POST"])
def api_export_md():
    sid=(request.json or {}).get("session_id"); conv=conversations.get(sid)
    if not conv: return jsonify({"error":"no conversation"}),404
    md=_conv_md(conv)
    return send_file(io.BytesIO(md.encode("utf-8")), mimetype="text/markdown",
                     as_attachment=True, download_name=(conv.get("title","墨核AI创作") or "墨核AI创作")+".md")
@app.route("/api/export/pdf", methods=["POST"])
def api_export_pdf():
    sid=(request.json or {}).get("session_id"); conv=conversations.get(sid)
    if not conv: return jsonify({"error":"no conversation"}),404
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    buf=io.BytesIO(); cv=canvas.Canvas(buf, pagesize=A4); cv.setFont("Helvetica",12); y=A4[1]-30
    for m in conv.get("messages",[]):
        role="你" if m["role"]=="user" else "墨核 AI"
        body=f"[{role}] "+m["content"]
        for line in body.split("\n"):
            for seg in [line[i:i+70] for i in range(0,len(line),70)]:
                cv.drawString(20, y, seg); y-=15
                if y<30: cv.showPage(); cv.setFont("Helvetica",12); y=A4[1]-30
        y-=10
    cv.save(); buf.seek(0)
    return send_file(buf, mimetype="application/pdf", as_attachment=True,
                     download_name=(conv.get("title","墨核AI创作") or "墨核AI创作")+".pdf")
@app.route("/api/export/wechat", methods=["POST"])
def api_export_wechat():
    sid=(request.json or {}).get("session_id"); conv=conversations.get(sid)
    if not conv: return jsonify({"error":"no conversation"}),404
    html=f"<h1>{conv.get('title','墨核 AI 创作')}</h1>"+ "".join(
        f"<p><b>{'你' if m['role']=='user' else '墨核 AI'}：</b><br>{m['content'].replace(chr(10),'<br>')}</p>" for m in conv.get("messages",[]))
    return send_file(io.BytesIO(html.encode("utf-8")), mimetype="text/html",
                     as_attachment=True, download_name=(conv.get("title","墨核AI创作") or "墨核AI创作")+".html")

@app.route("/api/export/xmind", methods=["POST"])
def api_export_xmind():
    """把对话导出为 XMind 思维导图（现代 .xmind = zip，内含 content.json）。"""
    sid=(request.json or {}).get("session_id"); conv=conversations.get(sid)
    if not conv: return jsonify({"error":"no conversation"}),404
    title=(conv.get("title") or "墨核AI创作")
    msgs=conv.get("messages",[])
    children=[]
    i=0
    while i < len(msgs):
        u=msgs[i]
        a=msgs[i+1] if i+1 < len(msgs) else None
        i += 2
        if u.get("role")!="user": continue
        node={"title": (u.get("content","") or "")[:60] or "提问"}
        if a and a.get("role")=="assistant":
            node["children"]={"attached":[{"title":"墨核 AI",
                "children":{"attached":[{"title":(a.get("content","") or "")[:400] or "…"}]}}]}
        children.append(node)
    root={"title": title}
    if children: root["children"]={"attached":children}
    content=[{"id":"sheet1","class":"sheet","rootTopic":root}]
    manifest={"file-entries":{"content.json":"","metadata.json":""}}
    metadata={"creator":{"name":"墨核AIStudio"},"dataStructureVersion":1}
    buf=io.BytesIO()
    with zipfile.ZipFile(buf,"w",zipfile.ZIP_DEFLATED) as z:
        z.writestr("content.json", json.dumps(content, ensure_ascii=False))
        z.writestr("metadata.json", json.dumps(metadata, ensure_ascii=False))
        z.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False))
    buf.seek(0)
    safe=(re.sub(r"[^\w一-龥-]","",title) or "墨核AI创作")
    return send_file(buf, mimetype="application/octet-stream", as_attachment=True, download_name=safe+".xmind")

# ---------- Markdown 编辑器（深度集成：自研 UI + 标准库导出，零侵权） ----------
import zipfile as _zipfile  # noqa  (前置声明，避免重复 import)
NOTES_DIR = DATA_HOME / "notes"
NOTES_DIR.mkdir(parents=True, exist_ok=True)

def _safe_note(name):
    """把任意输入收敛成 ~/.MoHeAI/notes 下的安全文件名，禁止路径穿越。"""
    name = (name or "").strip()
    if not name or "/" in name or "\\" in name or ".." in name or name.startswith("."):
        return None
    if not name.endswith(".md"):
        name += ".md"
    return NOTES_DIR / name

@app.route("/md")
def md_editor():
    """Markdown 编辑器整页（独立路由，主界面按钮跳转至此）。"""
    return send_from_directory(str(STATIC), "md/editor.html")

@app.route("/api/md/status")
def api_md_status():
    ai = _llm_usable(get_llm_cfg("md"))
    notes = [{"name": f.name[:-3], "mtime": int(f.stat().st_mtime)}
             for f in sorted(NOTES_DIR.glob("*.md"), key=lambda p: p.stat().st_mtime, reverse=True)]
    return jsonify({"ai_usable": ai, "notes": notes})

@app.route("/api/notes", methods=["GET", "POST", "DELETE"])
def api_notes():
    if request.method == "DELETE":
        p = _safe_note(request.args.get("name", ""))
        if p and p.exists():
            p.unlink()
        return jsonify({"ok": True})
    if request.method == "POST":
        p = _safe_note((request.json or {}).get("name", "") or request.args.get("name", ""))
        if not p:
            return jsonify({"error": "bad name"}), 400
        p.write_text((request.json or {}).get("content", ""), encoding="utf-8")
        return jsonify({"ok": True, "name": p.name[:-3]})
    notes = [{"name": f.name[:-3], "mtime": int(f.stat().st_mtime)}
             for f in sorted(NOTES_DIR.glob("*.md"), key=lambda p: p.stat().st_mtime, reverse=True)]
    return jsonify(notes)

@app.route("/api/notes/<name>")
def api_note_get(name):
    p = _safe_note(name)
    if not p or not p.exists():
        return jsonify({"error": "not found"}), 404
    return jsonify({"name": name, "content": p.read_text(encoding="utf-8")})

@app.route("/api/md/ai", methods=["POST"])
def api_md_ai():
    """选中文字 → 墨核 AI 润色 / 总结 / 翻译 / 续写 / 扩充（复用统一 LLM 调度）。"""
    d = request.json or {}
    text = (d.get("text") or "").strip()
    action = d.get("action", "polish")
    if not text:
        return jsonify({"error": "empty"}), 400
    if not _llm_usable(get_llm_cfg("md")):
        return jsonify({"result": "[墨核 AI 未配置] 请先在「设置」里启用大模型或本地内置模型，才能用 AI 润色 / 总结 / 翻译。"})
    SYS = {
        "polish":    "你是一位严谨的中文编辑。请在不改变原意的前提下，润色下面这段 Markdown 文本，使其更通顺、专业、连贯。只输出润色后的文本，不要解释。",
        "summarize": "请用简洁的中文总结下面这段 Markdown 文本的要点，用 Markdown 无序列表呈现。只输出总结。",
        "translate": "请将下面这段 Markdown 文本翻译成流畅的中文（若原文已是中文则英译中并保持）。保留 Markdown 格式。只输出译文。",
        "continue":  "请顺着下面这段 Markdown 文本继续写下去，保持原有风格与主题，续写一段内容。只输出新增的部分。",
        "expand":    "请在不改变原意的前提下，扩充下面这段 Markdown 文本，补充细节与例子，使其更丰满。只输出扩充后的全文。",
    }
    system = SYS.get(action, SYS["polish"])
    out = _llm_complete([{"role": "user", "content": system + "\n\n" + text}], get_llm_cfg("md"))
    return jsonify({"result": out})

@app.route("/api/md/export", methods=["POST"])
def api_md_export():
    """通用内容导出：txt / md / html / docx / xlsx / pdf（docx/xlsx 用标准库级合规库，零侵权）。"""
    d = request.json or {}
    content = d.get("content", "") or ""
    fmt = (d.get("fmt") or "md").lower()
    title = (d.get("title") or "墨核文档").strip() or "墨核文档"
    safe = re.sub(r"[^\w一-龥-]", "", title) or "墨核文档"
    if fmt in ("md", "txt"):
        return send_file(io.BytesIO(content.encode("utf-8")),
                         mimetype="text/markdown" if fmt == "md" else "text/plain;charset=utf-8",
                         as_attachment=True, download_name=safe + "." + fmt)
    if fmt == "html":
        try:
            import markdown as _md
            body = _md.markdown(content, extensions=["tables", "fenced_code"])
        except Exception:
            body = "<pre>" + content.replace("<", "&lt;") + "</pre>"
        html = ("<!DOCTYPE html><html lang=zh-CN><head><meta charset=utf-8><title>%s</title>"
                "<link rel=stylesheet href='/static/vendor/katex/katex.min.css'>"
                "<style>body{font-family:-apple-system,Segoe UI,'PingFang SC',sans-serif;max-width:840px;"
                "margin:40px auto;padding:0 20px;line-height:1.8;color:#222}"
                "code{background:#f2f2f2;padding:2px 6px;border-radius:4px}"
                "pre{background:#f6f8fa;padding:14px;border-radius:8px;overflow:auto}"
                "table{border-collapse:collapse;margin:10px 0}td,th{border:1px solid #ddd;padding:6px 10px}"
                "blockquote{color:#666;border-left:4px solid #ddd;margin:0;padding-left:14px}</style></head>"
                "<body>%s</body></html>") % (title, body)
        return send_file(io.BytesIO(html.encode("utf-8")), mimetype="text/html",
                         as_attachment=True, download_name=safe + ".html")
    if fmt == "docx":
        from docx import Document
        doc = Document()
        doc.add_heading(title, level=0)
        md_to_docx(doc, content)
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True, download_name=safe + ".docx")
    if fmt == "xlsx":
        import openpyxl
        from openpyxl import Workbook
        wb = Workbook(); ws = wb.active; ws.title = "全文"
        rows = [ln for ln in content.split("\n")]
        for i, ln in enumerate(rows, 1):
            ws.cell(row=i, column=1, value=ln)
        ws.column_dimensions["A"].width = 100
        # 解析 Markdown 表格 → 独立工作表
        lines = content.split("\n"); tbl_idx = 0; i = 0
        while i < len(lines):
            if lines[i].strip().startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
                header = [c.strip() for c in lines[i + 0].strip().strip("|").split("|")]
                tbl_idx += 1; ts = wb.create_sheet(title=("表格%d" % tbl_idx)[:31])
                for c, h in enumerate(header, 1):
                    ts.cell(row=1, column=c, value=h)
                r = 2
                j = i + 2
                while j < len(lines) and lines[j].strip().startswith("|"):
                    cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                    for c, v in enumerate(cells, 1):
                        ts.cell(row=r, column=c, value=v)
                    r += 1; j += 1
                i = j
            else:
                i += 1
        buf = io.BytesIO(); wb.save(buf); buf.seek(0)
        return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                         as_attachment=True, download_name=safe + ".xlsx")
    if fmt == "pdf":
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            cjk = "STSong-Light"
        except Exception:
            cjk = "Helvetica"
        buf = io.BytesIO(); cv = canvas.Canvas(buf, pagesize=A4)
        W, H = A4; y = H - 40; cv.setFont(cjk, 12)
        for raw in content.split("\n"):
            line = raw.rstrip()
            size = 12
            if line.startswith("### "): line, size = line[4:], 14
            elif line.startswith("## "): line, size = line[3:], 16
            elif line.startswith("# "): line, size = line[2:], 18
            cv.setFont(cjk, size)
            # 简单按宽度折行（CJK 约 1 字 ≈ size 宽）
            maxc = max(8, int((W - 80) / (size * 0.62)))
            while line:
                seg, line = line[:maxc], line[maxc:]
                cv.drawString(40, y, seg); y -= size + 6
                if y < 40:
                    cv.showPage(); cv.setFont(cjk, 12); y = H - 40
        cv.save(); buf.seek(0)
        return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name=safe + ".pdf")
    return jsonify({"error": "unsupported fmt: %s" % fmt}), 400

# ---------- Skill 导入 / 导出 ----------
@app.route("/api/skills/export", methods=["GET"])
def api_skills_export():
    return send_file(io.BytesIO(json.dumps(skills, ensure_ascii=False, indent=2).encode("utf-8")),
                     mimetype="application/json", as_attachment=True, download_name="skills.json")
@app.route("/api/skills/import", methods=["POST"])
def api_skills_import():
    global skills
    data=request.json or []
    if not isinstance(data, list): return jsonify({"error":"bad format"}),400
    skills.extend(data); _save(profile_dir()/"skills.json", skills)
    return jsonify({"ok":True,"count":len(skills)})
@app.route("/api/skills/gallery", methods=["GET"])
def api_skills_gallery():
    """技能市场：返回可一键添加的精选技能清单（不含已内置默认）。"""
    installed={s["id"] for s in skills}
    return jsonify([dict(g, installed=(g["id"] in installed)) for g in SKILL_GALLERY])
@app.route("/api/skills/import_url", methods=["POST"])
def api_skills_import_url():
    """从 URL 拉取社区分享的技能 JSON 并导入（Skill 市场/社区互通）。"""
    global skills
    url=(request.json or {}).get("url","")
    if not url: return jsonify({"error":"no url"}),400
    try:
        r=requests.get(url, timeout=15, headers={"User-Agent":"InkCore"})
        r.raise_for_status(); data=r.json()
    except Exception as e:
        return jsonify({"error":f"拉取失败：{e}"}),400
    if isinstance(data, dict): data=[data]
    if not isinstance(data, list): return jsonify({"error":"bad format"}),400
    skills.extend(data); _save(profile_dir()/"skills.json", skills)
    return jsonify({"ok":True,"count":len(skills),"added":len(data)})

# ---------- 离谱玩法：用户自建 / 导入导出（改动4） ----------
@app.route("/api/funs", methods=["GET","POST","DELETE"])
def api_funs():
    global funs
    if request.method=="GET": return jsonify(funs)
    if request.method=="POST":
        f=request.json or {}
        f["id"]=f.get("id") or f"fun_{uuid.uuid4().hex[:8]}"
        funs.append(f); _save(profile_dir()/"funs.json", funs); return jsonify(f)
    fid=request.args.get("id"); funs=[x for x in funs if x["id"]!=fid]; _save(profile_dir()/"funs.json", funs)
    return jsonify({"ok":True})
@app.route("/api/funs/export", methods=["GET"])
def api_funs_export():
    return send_file(io.BytesIO(json.dumps(funs, ensure_ascii=False, indent=2).encode("utf-8")),
                     mimetype="application/json", as_attachment=True, download_name="funs.json")
@app.route("/api/funs/import", methods=["POST"])
def api_funs_import():
    global funs
    data=request.json or []
    if not isinstance(data, list): return jsonify({"error":"bad format"}),400
    funs.extend(data); _save(profile_dir()/"funs.json", funs)
    return jsonify({"ok":True,"count":len(funs)})

# ---------- 版本与自动更新 ----------
@app.route("/api/version", methods=["GET"])
def api_version():
    return jsonify({"version":APP_VERSION})
@app.route("/api/check_update", methods=["POST"])
def api_check_update():
    """拉取远端 version.json 与本地版本比对。无网络/无清单时优雅返回。"""
    url=(request.json or {}).get("url") or UPDATE_URL_DEFAULT
    try:
        r=requests.get(url, timeout=15, headers={"User-Agent":"InkCore"})
        r.raise_for_status(); info=r.json()
        latest=str(info.get("version","")); notes=info.get("notes","")
        dl=info.get("url","")
        def _vp(v):
            return tuple(int(x) for x in re.findall(r"\d+", v)[:3])
        avail = bool(latest) and _vp(latest) > _vp(APP_VERSION)
        return jsonify({"current":APP_VERSION,"latest":latest,"update_available":avail,
                        "download_url":dl,"notes":notes})
    except Exception as e:
        return jsonify({"current":APP_VERSION,"latest":None,"update_available":False,
                        "error":f"检查失败（可能无外网）：{e}"})

# ---------- Agent 运行（服务端真实编排） ----------
@app.route("/api/agent", methods=["POST"])
def api_agent():
    d=request.json or {}; goal=d.get("goal",""); steps=d.get("steps",[])
    if not steps: return jsonify({"error":"empty steps"}),400
    return jsonify({"reply":agent_run(goal, steps),"mode":"agent"})

# ---------- 本地内置模型（离线推理，免 Ollama） ----------
@app.route("/api/embedded/status")
def api_embedded_status():
    """返回本地模型的下载 / 加载状态（前端轮询）。"""
    return jsonify(local_llm.get_status())

@app.route("/api/embedded/download", methods=["POST"])
def api_embedded_download():
    """后台启动模型权重下载（首次使用触发）；立即返回，前端轮询 status。"""
    d = request.json or {}
    key = d.get("model") or None
    ok = local_llm.start_download(key)
    if not ok:
        return jsonify({"error": "已在下载中"}), 409
    return jsonify({"started": True, "model": local_llm._default_model_key(key)})

# ---------- 数据备份 / 恢复 ----------
@app.route("/api/backup")
def api_backup():
    """打包整个 ~/MoHeAI 数据目录为 zip 下载。"""
    buf=io.BytesIO()
    with zipfile.ZipFile(buf,"w",zipfile.ZIP_DEFLATED) as z:
        for p in sorted(DATA_HOME.rglob("*")):
            if p.is_file():
                z.write(p, str(p.relative_to(DATA_HOME)))
    buf.seek(0)
    return send_file(buf, mimetype="application/zip", as_attachment=True, download_name="MoHeAI_backup.zip")

@app.route("/api/restore", methods=["POST"])
def api_restore():
    """从备份 zip 恢复（仅解压进 ~/MoHeAI，校验路径防止越权）。"""
    f=request.files.get("file")
    if not f: return jsonify({"error":"no file"}),400
    tmp=DATA_HOME/"_restore_tmp.zip"; f.save(str(tmp))
    try:
        with zipfile.ZipFile(str(tmp)) as z:
            for name in z.namelist():
                if ".." in name or name.startswith("/") or name.startswith("\\"):
                    return jsonify({"error":"非法路径，已拒绝"}),400
            z.extractall(str(DATA_HOME))
        # 重新加载当前 profile 内存态，避免旧数据残留
        load_profile()
        return jsonify({"ok":True})
    except Exception as e:
        return jsonify({"error":str(e)}),400
    finally:
        try: tmp.unlink()
        except Exception: pass

if __name__ == "__main__":
    print("墨核 AI Studio (InkCore) 已启动 → http://127.0.0.1:7860")
    try:
        # 生产级 WSGI 服务器：支持并发 SSE 流式 + 普通请求（dev server 会在并发流时卡住）
        from waitress import serve
        serve(app, host="127.0.0.1", port=7860, threads=12, channel_timeout=120)
    except Exception as e:
        print("waitress 不可用，回退到 Flask dev server：", e)
        app.run(host="127.0.0.1", port=7860, debug=False, threaded=True)
